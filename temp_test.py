import os
import re
import tkinter as tk
from tkinter import filedialog
import docx
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image  # 需先 pip install Pillow
from docx2pdf import convert as docx2pdf_convert
from PyPDF2 import PdfMerger
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_top_border_bold(cell, border_color="000000", border_size="6", border_space="0", border_val="single"):
    """
    將指定儲存格的上邊框設為 3/4pt（即 border_size 為 "6"）
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 找出現有的邊框設定
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    else:
        # 移除現有的上邊框設定（如果有的話）
        for top in tcBorders.findall(qn('w:top')):
            tcBorders.remove(top)
    # 新增上邊框元素
    top = OxmlElement('w:top')
    top.set(qn('w:val'), border_val)
    top.set(qn('w:sz'), border_size)
    top.set(qn('w:space'), border_space)
    top.set(qn('w:color'), border_color)
    tcBorders.append(top)

def set_cell_font(cell, font_name="標楷體"):
    """
    將指定儲存格所有段落中的文字字型設為 font_name
    """
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            # 設定東亞文字字型
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), font_name)

def get_image_files(folder):
    """
    取得指定資料夾中所有圖片檔案（依副檔名過濾）。
    """
    valid_extensions = [".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff"]
    return [
        os.path.join(folder, f)
        for f in os.listdir(folder)
        if os.path.splitext(f)[1].lower() in valid_extensions
    ]

def generate_dummy_image(filename):
    """
    產生一張全白的圖片並儲存，回傳該圖片的檔案路徑。
    這裡設定圖片大小為 500x500，可依需求調整。
    """
    img = Image.new("RGB", (500, 500), "white")
    dummy_path = os.path.join(os.getcwd(), filename)
    img.save(dummy_path)
    return dummy_path

def insert_images_into_9x3_template_left_to_right(template_path, images, output_prefix):
    """
    使用 9×3 的 Word 模板（第一行與第一列保留），
    有效區域為第 2～9 行、列 2～3，共 8 個儲存格區塊（分為 4 個 row pair × 2 欄）。

    填入規則（依 row-major order）：
      - 對於圖片 index j (0-based)：
          block = j (共 0~7，每份文件最多 8 張圖片)
          row_pair = block // 2  (0～3)
          effective col index = block % 2  (0 或 1)
          絕對欄 = 1 + effective col index (即 1 或 2)
          絕對檔名行 = 1 + 2 * row_pair  (分別為 1, 3, 5, 7)
          絕對圖片行 = 1 + 2 * row_pair + 1 (分別為 2, 4, 6, 8)
      - 檔名儲存格：若檔名以 blank 開頭則顯示空字串，
        否則取檔名去除副檔名並在最前面加上「編號:」；
        對於「埋深照」和「銑鋪照」的檔名，若檔名僅由數字組成，則去除前導零；
        對於「讀數照」的檔名，若以 "app_" 為前綴則去除該前綴。
      - 當該圖片為某類別（category）的第一張時，
        將左側儲存格（第 0 欄）與下方儲存格合併，
        並在合併後的儲存格中填入預先設定好的垂直文字（每字換行），
        且儲存格上方邊框加粗，文字水平置中。
      - 在圖片儲存格插入圖片（寬度設定為 5 公分）。

    在每份 Word 文件完成前，額外檢查第二行第一列（table.cell(1,0)）是否有資料，
    如果沒有，則根據該組中第二個圖片的類別，從 vertical_text_dict 取得對應的文字填入該儲存格。

    每組 8 張圖片產生一份文件，檔名格式： temp_{base}_{組號}.docx
    回傳生成的檔案路徑列表。
    """
    output_files = []
    groups = [images[i:i+8] for i in range(0, len(images), 8)]
    
    out_dir = os.path.dirname(output_prefix)
    base_name = os.path.basename(output_prefix)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir)
    
    # 用來記錄已插入類別名稱的種類（避免重複插入）
    used_categories = set()
    
    # 更新後的垂直文字字典
    vertical_text_dict = {
        "埋深照": "二\n、\n深\n度\n相\n片",
        "銑鋪照": "三\n、\n臨\n時\n修\n復\n後\n全\n景\n照\n片",
        "測量照": "四\n、\n施\n測\n相\n片",
        "讀數照": "五\n、\n讀\n數\n相\n片"
    }
    
    for group_idx, group in enumerate(groups, start=1):
        doc = docx.Document(template_path)
        table = doc.tables[0]  # 假設模板中的第一個表格為 9×3
        
        for j, (img_path, category) in enumerate(group):
            row_pair = j // 2            # 0～3
            effective_col = j % 2        # 0 或 1
            abs_col = 1 + effective_col  # 有效欄位（1 或 2）
            abs_row_fname = 1 + 2 * row_pair      # 檔名行：1, 3, 5, 7
            abs_row_img   = 1 + 2 * row_pair + 1  # 圖片行：2, 4, 6, 8

            # 檔名儲存格：先移除副檔名，再加上"編號:"前綴；若檔名以 blank 開頭則留空
            fname_cell = table.cell(abs_row_fname, abs_col)
            basename = os.path.basename(img_path)
            if basename.startswith("blank"):
                fname_cell.text = ""
            else:
                name_no_ext = os.path.splitext(basename)[0]
                # 對於埋深照和銑鋪照，若檔名僅由數字組成，則去除前導零
                if category in ["埋深照", "銑鋪照"]:
                    try:
                        name_no_ext = str(int(name_no_ext))
                    except:
                        pass
                # 對於讀數照，若檔名以 "app_" 為前綴則去除
                if category == "讀數照" and name_no_ext.startswith("app_"):
                    name_no_ext = name_no_ext[4:]
                fname_cell.text = "編號:" + name_no_ext
            for para in fname_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(fname_cell, "標楷體")
            
            # 當該圖片為該類別的第一張時，
            # 合併左側儲存格（第 0 欄）與下方儲存格，
            # 並插入預先設定好的垂直文字，文字水平置中，並將上邊框加粗
            if category not in used_categories:
                merged_cell = table.cell(abs_row_fname, 0).merge(table.cell(abs_row_img, 0))
                vertical_text = vertical_text_dict.get(category, category)
                merged_cell.text = vertical_text
                for para in merged_cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_top_border_bold(merged_cell)
                set_cell_font(merged_cell, "標楷體")
                used_categories.add(category)
            
            # 插入圖片至圖片儲存格
            img_cell = table.cell(abs_row_img, abs_col)
            img_cell.text = ""
            paragraph = img_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(img_cell, "標楷體")
            run = paragraph.add_run()
            run.add_picture(img_path, height=Cm(5.47))
        
        # 檢查第二行第一列的儲存格（table.cell(1,0)）是否有資料，
        # 如果沒有，則取該組中第二個圖片的類別並填入對應的垂直文字
        if table.cell(1, 0).text.strip() == "" and len(group) > 1:
            category = group[1][1]
            vertical_text = vertical_text_dict.get(category, category)
            cell_to_fill = table.cell(1, 0)
            cell_to_fill.text = vertical_text
            for para in cell_to_fill.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_top_border_bold(cell_to_fill)
            set_cell_font(cell_to_fill, "標楷體")
        
        output_file = os.path.join(out_dir, f"temp_{base_name}_{group_idx}.docx")
        doc.save(output_file)
        print(f"已儲存：{output_file}")
        output_files.append(output_file)
    
    return output_files

def main():
    root = tk.Tk()
    root.withdraw()
    base_folder = filedialog.askdirectory(title="請選擇包含圖片的資料夾")
    if not base_folder:
        print("未選擇資料夾，程式結束。")
        return

    images = []
    # 資料夾路徑與類別名稱
    folder1 = os.path.join(base_folder, "埋深照")
    folder2 = os.path.join(base_folder, "銑鋪照")
    folder3 = os.path.join(base_folder, "測量照")
    folder4 = os.path.join(base_folder, "讀數照")

    # 處理埋深照
    if os.path.exists(folder1):
        images1 = get_image_files(folder1)
        if not images1:
            images1 = [generate_dummy_image("blank_0.jpg"), generate_dummy_image("blank_0_2.jpg")]
        elif len(images1) % 2 == 1:
            images1.append(generate_dummy_image("blank_0.jpg"))
    else:
        images1 = [generate_dummy_image("blank_0.jpg"), generate_dummy_image("blank_0_2.jpg")]
    images.extend([(img, "埋深照") for img in images1])

    # 處理銑鋪照
    if os.path.exists(folder2):
        images2 = get_image_files(folder2)
        if not images2:
            images2 = [generate_dummy_image("blank_1.jpg"), generate_dummy_image("blank_1_2.jpg")]
        elif len(images2) % 2 == 1:
            images2.append(generate_dummy_image("blank_1.jpg"))
    else:
        images2 = [generate_dummy_image("blank_1.jpg"), generate_dummy_image("blank_1_2.jpg")]
    images.extend([(img, "銑鋪照") for img in images2])

    # 處理測量照：依照檔名中的數字自然排序
    images3 = []
    if os.path.exists(folder3):
        files3 = get_image_files(folder3)
        sorted_files3 = sorted(files3, key=lambda x: int(re.search(r'\d+', os.path.basename(x)).group()))
        images3 = sorted_files3
    images.extend([(img, "測量照") for img in images3])

    # 處理讀數照：同上
    images4 = []
    if os.path.exists(folder4):
        files4 = get_image_files(folder4)
        sorted_files4 = sorted(files4, key=lambda x: int(re.search(r'\d+', os.path.basename(x)).group()))
        images4 = sorted_files4
    images.extend([(img, "讀數照") for img in images4])

    if not images:
        print("找不到任何圖片，程式結束。")
        return

    template_path = r"template\自主查核表_表格模板.docx"
    output_prefix = r"output\images_filled"
    word_files = insert_images_into_9x3_template_left_to_right(template_path, images, output_prefix)

    # 轉換所有 Word 檔案為 PDF
    pdf_files = []
    for word_file in word_files:
        pdf_file = word_file.replace(".docx", ".pdf")
        docx2pdf_convert(word_file, pdf_file)
        print(f"已轉換為 PDF：{pdf_file}")
        pdf_files.append(pdf_file)

    # 合併所有 PDF 為一份 PDF
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)
    merged_pdf_path = os.path.join(os.getcwd(), "merged_output.pdf")
    merger.write(merged_pdf_path)
    merger.close()
    print(f"已合併 PDF：{merged_pdf_path}")

if __name__ == "__main__":
    main()
