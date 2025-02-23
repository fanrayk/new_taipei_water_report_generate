# doc_generator.py
import os
from docxtpl import DocxTemplate
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx2pdf import convert
from PyPDF2 import PdfMerger
from utils import set_cell_width

def generate_records_doc(record, output_folder):
    template_path = os.path.join(
        "template", "自主查核表_首頁模板.docx"
    )
    doc = DocxTemplate(template_path)
    doc.render(record)
    # 在檔名前加上 context_number
    docx_filename = os.path.join(output_folder, "temp_自主查核表首頁.docx")
    doc.save(docx_filename)
    pdf_path = os.path.join(output_folder, "temp_自主查核表首頁.pdf")
    convert(docx_filename, pdf_path)
    print("Records PDF 已產生：", pdf_path)
    return pdf_path


def generate_pipeline_doc(simulated_data, context_number, output_folder):
    template_path = os.path.join(
        "template", "附件1模板", "附件1_定位資料回饋表_管道模板.docx"
    )
    doc = DocxTemplate(template_path)
    subdoc = doc.new_subdoc()
    num_cols = 7
    table = subdoc.add_table(rows=1, cols=num_cols)
    headers = ["編號", "種類", "座標X", "座標Y", "地盤高程", "埋管深度", "管頂座標z"]
    for i, cell in enumerate(table.rows[0].cells):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.first_line_indent = 0
        run = paragraph.add_run(headers[i])
        run.font.name = "標楷體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_row = table.rows[0]
    header_row.height = Pt(30)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for item in simulated_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["Number"])
        row_cells[1].text = str(item["Type"])
        row_cells[2].text = str(item["Coordinate_X"])
        row_cells[3].text = str(item["Coordinate_Y"])
        row_cells[4].text = str(item["Ground_Elevation"])
        row_cells[5].text = str(item["Pipe_Burial_Depth"])
        row_cells[6].text = str(item["Pipe_Top_Coordinate_Z"])
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = 0
                paragraph.paragraph_format.first_line_indent = 0
                for run in paragraph.runs:
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    tbl = table._element
    tblPr_list = tbl.xpath("./w:tblPr")
    tblPr = tblPr_list[0] if tblPr_list else OxmlElement("w:tblPr")
    if not tblPr_list:
        tbl.insert(0, tblPr)
    tblW_list = tblPr.xpath("./w:tblW")
    tblW = tblW_list[0] if tblW_list else OxmlElement("w:tblW")
    if not tblW_list:
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "10000")
    tblW.set(qn("w:type"), "dxa")
    column_widths = [658, 1756, 1316, 1429, 1094, 1094, 1208]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[idx])
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tbl_borders.append(border)
    tblPr.append(tbl_borders)
    context = {"table": subdoc, "case_number": context_number}
    doc.render(context)
    word_filename = os.path.join(output_folder, "temp_管線.docx")
    doc.save(word_filename)
    pdf_filename = os.path.join(output_folder, "temp_管線.pdf")
    convert(word_filename, pdf_filename)
    print("管線 PDF 已產生：", pdf_filename)
    return pdf_filename


def generate_reserved_doc(reserved_data, context_number, output_folder):
    template_path = os.path.join(
        "template", "附件1模板", "附件1_定位資料回饋表_設施物模板.docx"
    )
    doc = DocxTemplate(template_path)
    subdoc = doc.new_subdoc()
    num_cols = 6
    table = subdoc.add_table(rows=1, cols=num_cols)
    reserved_headers = [
        "編號",
        "種類",
        "座標X",
        "座標Y",
        "地盤高程",
        "座標z",
    ]
    for i, cell in enumerate(table.rows[0].cells):
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.left_indent = 0
        paragraph.paragraph_format.first_line_indent = 0
        run = paragraph.add_run(reserved_headers[i])
        run.font.name = "標楷體"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_row = table.rows[0]
    header_row.height = Pt(30)
    header_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for item in reserved_data:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["Number"])
        row_cells[1].text = str(item["Type"])
        row_cells[2].text = str(item["Coordinate_X"])
        row_cells[3].text = str(item["Coordinate_Y"])
        row_cells[4].text = str(item["Ground_Elevation"])
        row_cells[5].text = str(item["Pipe_Top_Coordinate_Z"])
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.left_indent = 0
                paragraph.paragraph_format.first_line_indent = 0
                for run in paragraph.runs:
                    run.font.name = "標楷體"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")
    tbl = table._element
    tblPr_list = tbl.xpath("./w:tblPr")
    tblPr = tblPr_list[0] if tblPr_list else OxmlElement("w:tblPr")
    if not tblPr_list:
        tbl.insert(0, tblPr)
    tblW_list = tblPr.xpath("./w:tblW")
    tblW = tblW_list[0] if tblW_list else OxmlElement("w:tblW")
    if not tblW_list:
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "10000")
    tblW.set(qn("w:type"), "dxa")
    column_widths = [658, 1756, 1316, 1429, 1094, 1208]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, column_widths[idx])
    tbl_borders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement("w:" + border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tbl_borders.append(border)
    tblPr.append(tbl_borders)
    context = {"table": subdoc, "cast_number": context_number}
    doc.render(context)
    docx_filename = os.path.join(output_folder, "temp_設施物.docx")
    doc.save(docx_filename)
    pdf_filename = os.path.join(output_folder, "temp_設施物.pdf")
    convert(docx_filename, pdf_filename)
    print("設施物 PDF 已產生：", pdf_filename)
    return pdf_filename


def merge_pdf_files(pdf_files, merged_pdf_filename):
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)
    merger.write(merged_pdf_filename)
    merger.close()
    print("PDF 合併完成，最終 PDF 檔案：", merged_pdf_filename)
