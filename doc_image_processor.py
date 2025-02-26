import os
import docx
from docx.shared import Cm
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PyPDF2 import PdfMerger
from typing import List

def set_vertical_text_alternative(cell: docx.table._Cell, text: str) -> None:
    """
    將 cell 內容設定為指定文字，每個字換行以達到垂直排列效果，
    並設定字型為標楷體。
    """
    cell.text = ""
    vertical_text = "\n".join(list(text))
    paragraph = cell.add_paragraph(vertical_text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = 0
    paragraph.paragraph_format.first_line_indent = 0
    run = paragraph.runs[0]
    run.font.name = "標楷體"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")


def get_image_files(plane_folder: str) -> List[str]:
    """
    取得指定資料夾中所有圖片檔案（依副檔名過濾）。
    """
    valid_extensions = [".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff"]
    return [os.path.join(plane_folder, f)
            for f in os.listdir(plane_folder)
            if os.path.splitext(f)[1].lower() in valid_extensions]


def insert_images_in_template(template_path: str, image_group: List[str], output_file: str) -> None:
    """
    根據模板將圖片與文字插入至表格指定區域：
      - 第一張圖片放在第2~5行、第2~3欄合併儲存格內（垂直置中）。
      - 第二張圖片放在第6~9行、第2~3欄合併儲存格內（垂直置中）。
      - 合併第一欄第2~9行儲存格，並以垂直排列方式插入固定文字。
      
    產生的 Word 文件將儲存至 output_file。
    """
    doc = docx.Document(template_path)
    table = doc.tables[0]

    # 第一張圖片（若存在）
    if len(image_group) >= 1:
        merged_cell = table.cell(1, 1).merge(table.cell(4, 2))
        merged_cell.text = ""
        paragraph = merged_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_group[0], width=Cm(15.91))
        merged_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 第二張圖片（若存在）
    if len(image_group) >= 2:
        merged_cell2 = table.cell(5, 1).merge(table.cell(8, 2))
        merged_cell2.text = ""
        paragraph2 = merged_cell2.paragraphs[0]
        run2 = paragraph2.add_run()
        run2.add_picture(image_group[1], width=Cm(15.91))
        merged_cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    # 合併第一欄第2~9行，插入垂直排列文字
    merged_text_cell = table.cell(1, 0).merge(table.cell(8, 0))
    set_vertical_text_alternative(merged_text_cell, "一、竣工平面圖")

    doc.save(output_file)
    print(f"已儲存 Word 文件：{output_file}")


def convert_word_to_pdf(word_path: str, pdf_path: str) -> None:
    """
    將指定的 Word 文件轉換成 PDF。
    """
    try:
        convert(word_path, pdf_path)
        print(f"已生成 PDF 文件：{pdf_path}")
    except AttributeError as e:
        if "Word.Application.Quit" in str(e):
            print(f"遇到 Word.Application.Quit 錯誤於 {word_path}，忽略並繼續。")
        else:
            raise e


def merge_pdfs(pdf_folder: str, output_pdf: str) -> None:
    """
    將 pdf_folder 中所有檔名符合 temp_modified_template_group_*.pdf 的 PDF 檔案合併，
    並儲存為 output_pdf。
    """
    pdf_files = [os.path.join(pdf_folder, f)
                 for f in os.listdir(pdf_folder)
                 if f.endswith('.pdf') and f.startswith("temp_modified_template_group_")]
    pdf_files.sort()
    if pdf_files:
        merger = PdfMerger()
        for pdf in pdf_files:
            merger.append(pdf)
        merger.write(output_pdf)
        merger.close()
        print(f"已合併所有 PDF 至：{output_pdf}")


def merge_pdfs_from_list(pdf_list: List[str], output_pdf: str) -> None:
    """
    合併 pdf_list 中所有 PDF 檔案，並輸出至 output_pdf。
    """
    merger = PdfMerger()
    for pdf in pdf_list:
        merger.append(pdf)
    merger.write(output_pdf)
    merger.close()
    print(f"合併完成：{output_pdf}")


def process_documents(main_folder: str, template_path: str, output_folder: str, case_number: str) -> None:
    """
    核心流程：
      1. 從 main_folder 下的「平面圖」資料夾取得所有圖片。
      2. 每兩張圖片一組，根據模板產生對應的 Word 文件並轉換為 PDF。
      3. 合併所有 PDF 至一份。
    """
    plane_folder = os.path.join(main_folder, "平面圖")
    if not os.path.isdir(plane_folder):
        print(f"找不到資料夾：{plane_folder}")
        return

    images = get_image_files(plane_folder)
    print(f"平面圖資料夾中共有 {len(images)} 張圖片。")
    if not images:
        print("沒有找到任何圖片，程式終止。")
        return

    groups = [images[i:i+2] for i in range(0, len(images), 2)]
    print(f"總共分成 {len(groups)} 組。")

    for idx, group in enumerate(groups, start=1):
        word_filename = f"temp_modified_template_group_{idx}.docx"
        word_path = os.path.join(output_folder, word_filename)
        pdf_filename = f"temp_modified_template_group_{idx}.pdf"
        pdf_path = os.path.join(output_folder, pdf_filename)

        insert_images_in_template(template_path, group, word_path)
        convert_word_to_pdf(word_path, pdf_path)

    merged_pdf_path = os.path.join(output_folder, f"temp_{case_number}_平面圖.pdf")
    merge_pdfs(output_folder, merged_pdf_path)


# 以下為從原 main_helpers.py 移入的 docx 相關函式

def set_cell_top_border_bold(cell, border_color="000000", border_size="6", border_space="0", border_val="single") -> None:
    """
    將指定儲存格的上邊框設為 3/4pt（即 border_size 為 "6"）。
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    else:
        # 移除現有的上邊框設定
        for top in tcBorders.findall(qn("w:top")):
            tcBorders.remove(top)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), border_val)
    top.set(qn("w:sz"), border_size)
    top.set(qn("w:space"), border_space)
    top.set(qn("w:color"), border_color)
    tcBorders.append(top)


def set_cell_font(cell, font_name="標楷體") -> None:
    """
    將指定儲存格所有段落中的文字字型設為 font_name。
    """
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def insert_images_into_9x3_template_left_to_right(template_path: str, images: list, output_prefix: str) -> list:
    """
    使用 9×3 Word 模板將圖片依規則填入指定儲存格中，
    每份文件最多 8 張圖片，並依據圖片類別插入標題文字。
    
    :param template_path: 模板檔案路徑
    :param images: [(圖片路徑, 類別)] 的列表
    :param output_prefix: 輸出檔案的前置名稱
    :return: 生成的 Word 檔案路徑列表
    """
    output_files = []
    groups = [images[i: i+8] for i in range(0, len(images), 8)]
    out_dir = os.path.dirname(output_prefix)
    base_name = os.path.basename(output_prefix)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir)

    used_categories = set()
    vertical_text_dict = {
        "埋深照": "二\n、\n深\n度\n相\n片",
        "銑鋪照": "三\n、\n臨\n時\n修\n復\n後\n全\n景\n照\n片",
        "測量照": "四\n、\n施\n測\n相\n片",
        "讀數照": "五\n、\n讀\n數\n相\n片",
    }

    # 為避免重複引入，內部局部引用
    import re
    import docx
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for group_idx, group in enumerate(groups, start=1):
        doc = docx.Document(template_path)
        table = doc.tables[0]

        for j, (img_path, category) in enumerate(group):
            row_pair = j // 2
            effective_col = j % 2
            abs_col = 1 + effective_col
            abs_row_fname = 1 + 2 * row_pair
            abs_row_img = 1 + 2 * row_pair + 1

            fname_cell = table.cell(abs_row_fname, abs_col)
            basename = os.path.basename(img_path)
            if basename.startswith("blank"):
                fname_cell.text = ""
            else:
                name_no_ext = os.path.splitext(basename)[0]
                if category in ["埋深照", "銑鋪照"]:
                    try:
                        name_no_ext = str(int(name_no_ext))
                    except Exception:
                        pass
                if category == "讀數照" and name_no_ext.startswith("app_"):
                    name_no_ext = name_no_ext[4:]
                fname_cell.text = "編號:" + name_no_ext
            for para in fname_cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(fname_cell, "標楷體")

            if category not in used_categories:
                merged_cell = table.cell(abs_row_fname, 0).merge(table.cell(abs_row_img, 0))
                vertical_text = vertical_text_dict.get(category, category)
                merged_cell.text = vertical_text
                for para in merged_cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_top_border_bold(merged_cell)
                set_cell_font(merged_cell, "標楷體")
                used_categories.add(category)

            img_cell = table.cell(abs_row_img, abs_col)
            img_cell.text = ""
            paragraph = img_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_font(img_cell, "標楷體")
            run = paragraph.add_run()
            run.add_picture(img_path, height=Cm(5.47))

        if table.cell(1, 0).text.strip() == "" and len(group) > 1:
            category = group[1][1]
            vertical_text = vertical_text_dict.get(category, category)
            cell_to_fill = table.cell(1, 0)
            cell_to_fill.text = vertical_text
            for para in cell_to_fill.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_top_border_bold(cell_to_fill)
            set_cell_font(cell_to_fill, "標楷體")

        output_file = os.path.join(out_dir, f"temp_{base_name}_{group_idx}.docx")
        doc.save(output_file)
        print(f"已儲存：{output_file}")
        output_files.append(output_file)

    return output_files
